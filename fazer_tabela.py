from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_ALIGN_VERTICAL
from tkinter import simpledialog, messagebox, ttk
import tkinter as tk
from tkinter.simpledialog import askstring
from servicos_proposta import final_services

def centralizar_prompt(mensagem, titulo=''):
    root = tk.Tk()
    root.withdraw()
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    prompt_width, prompt_height = 400, 100
    pos_x = (screen_width - prompt_width) // 2
    pos_y = (screen_height - prompt_height) // 2

    root.geometry(f'{prompt_width}x{prompt_height}+{pos_x}+{pos_y}')
    root.configure(bg='#2b2b2b')
    
    response = askstring(titulo, mensagem, parent=root)
    root.destroy()
    return response

def criar_tabela_word(cabecalhos, lista_descricoes, documento_meclar):
    documento1 = documento_meclar
    
    table = documento1.add_table(rows=1, cols=len(cabecalhos))
    table.style = 'Table Grid'

    column_widths = [Inches(1.5), Inches(3), Inches(1.5), Inches(1.5), Inches(2), Inches(2)]
    for col_idx, width in enumerate(column_widths):
        for cell in table.columns[col_idx].cells:
            cell.width = width

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(cabecalhos):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(12)
        hdr_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'  # Ajuste da fonte
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shading_elm = parse_xml(r'<w:shd {} w:fill="B7DEE8"/>'.format(nsdecls('w')))
        hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)

    table.rows[0].height = Inches(0.5)

    total_geral = 0  
    
    root = tk.Tk()
    root.title("Entrada de Dados")
    root.configure(bg='#3c3f41')
    
    style = ttk.Style(root)
    style.configure("TButton", font=('Helvetica', 12), padding=10)
    style.configure("TLabel", font=('Helvetica', 12))
    style.configure("TEntry", font=('Helvetica', 12), padding=5)

    entries = []

    # Criação do tooltip
    tooltip = tk.Toplevel(root)
    tooltip.wm_overrideredirect(True)  # Remove as bordas e a barra de título
    tooltip_label = tk.Label(tooltip, text="", background="#FFFFE0", relief='solid', borderwidth=1, wraplength=400)
    tooltip_label.pack()
    tooltip.withdraw()  # Oculta o tooltip inicialmente

    def show_tooltip(event, text):
        tooltip_label.config(text=text)
        tooltip.geometry(f"+{event.x_root + 10}+{event.y_root + 10}")
        tooltip.deiconify()  # Mostra o tooltip

    def hide_tooltip(event):
        tooltip.withdraw()  # Oculta o tooltip

    for i, descricao in enumerate(lista_descricoes):
        frame = ttk.Frame(root)
        frame.pack(fill='x', pady=5)
        
        # Ajustando o label para ocupar mais espaço
        label = ttk.Label(frame, text=descricao[:25] + '...' if len(descricao) > 25 else descricao, width=35, anchor='w')
        label.pack(side='left', padx=10, fill='x')
        
        quantidade_var = tk.StringVar(value='1')
        entry_quantidade = ttk.Entry(frame, textvariable=quantidade_var, width=10)
        entry_quantidade.pack(side='left', padx=5)
        
        custo_unitario_var = tk.StringVar(value='R$ 100.00')
        entry_custo = ttk.Entry(frame, textvariable=custo_unitario_var, width=10)
        entry_custo.pack(side='left', padx=5)
        
        label.bind("<Enter>", lambda e, text=descricao: show_tooltip(e, text))
        label.bind("<Leave>", hide_tooltip)
        
        entries.append((descricao, quantidade_var, custo_unitario_var))

    def on_calculate_total():
        nonlocal total_geral
        total_geral = 0  

        for descricao, quantidade_var, custo_unitario_var in entries:
            quantidade = int(quantidade_var.get())
            custo_unitario = float(custo_unitario_var.get().replace("R$", "").strip())
            custo_total = quantidade * custo_unitario
            total_geral += custo_total  

        total_ajustado_var.set(f"R$ {total_geral:.2f}")

    calculate_button = ttk.Button(root, text="Calcular Total", command=on_calculate_total)
    calculate_button.pack(pady=10)

    frame_total = ttk.Frame(root)
    frame_total.pack(fill='x', pady=10)
    
    ttk.Label(frame_total, text="Ajuste do Valor Total:", width=30).pack(side='left', padx=10)
    
    total_ajustado_var = tk.StringVar(value='R$ 0.00')
    ttk.Entry(frame_total, textvariable=total_ajustado_var, width=10).pack(side='left', padx=5)

    def on_submit():
        nonlocal total_geral
        for i, (descricao, quantidade_var, custo_unitario_var) in enumerate(entries):
            quantidade = int(quantidade_var.get())
            custo_unitario = float(custo_unitario_var.get().replace("R$", "").strip())
            unidade = "km" if descricao == "Deslocamento da equipe técnica" else "serviço"
            custo_total = quantidade * custo_unitario
            
            row_cells = table.add_row().cells
            row_cells[0].text = str(i + 1)
            row_cells[1].text = descricao
            row_cells[2].text = unidade
            row_cells[3].text = str(quantidade)
            row_cells[4].text = f"R$ {custo_unitario:.2f}"
            row_cells[5].text = f"R$ {custo_total:.2f}"

            for cell in row_cells:
                cell.paragraphs[0].runs[0].font.size = Pt(12)
                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'  # Ajuste da fonte
                
            table.rows[i + 1].height = Inches(0.4)
        
        row_cells = table.add_row().cells
        row_cells[0].text = ''
        row_cells[1].text = ''
        row_cells[2].text = ''
        row_cells[3].text = ''
        row_cells[4].text = 'TOTAL GERAL'
        row_cells[5].text = f"R$ {total_geral:.2f}"

        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'  # Ajuste da fonte
            
        table.rows[-1].height = Inches(0.4)
        
        total_ajustado = total_ajustado_var.get().replace("R$", "").strip()
        if total_ajustado and total_ajustado.lower() != 'ok':
            row_cells[5].text = f"R$ {float(total_ajustado):.2f}"

        root.destroy()

    ttk.Button(root, text="Enviar", command=on_submit).pack(pady=10)
    
    # Adiciona uma animação para o aparecimento da janela
    for i in range(0, 100, 5):
        root.attributes('-alpha', i/100)
        root.update()
        root.after(10)
    
    root.mainloop()

    return total_geral  
