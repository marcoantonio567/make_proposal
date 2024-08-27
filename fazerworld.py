from docx import Document
import tkinter as tk
from tkinter import simpledialog
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from funcoes import mes_atual , calcular_percentual , dia_atual , ano_atual ,formatar_cpf
from servicos_proposta import final_descriptions, final_services
from docx.enum.text import WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT 
from docx.enum.table import WD_ALIGN_VERTICAL 
from fazer_tabela import  criar_tabela_word  
import pyautogui
from cabecalho_proposta import *
import os


# Função para adicionar uma linha em branco
def add_blank_line(doc, size=12):
    p = doc.add_paragraph()
    run = p.add_run()
    run.font.size = Pt(size)
    run.add_break()

# Função para adicionar o rodapé em todas as páginas, exceto a primeira seção
def add_footer(section):
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer_paragraph.add_run(
        "Quadra 304 Norte, Avenida LO 08, Lote 01A, Sala 6, Plano Diretor Norte. "
        "CEP: 77.006-348, Palmas – TO, Fone: (63) 99221 – 6383, e-mail: atendimento@agropassos.eng.br "
        "www.agropassos.eng.br"
    )
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

# Criação do documento
doc = Document()

# Adicionando a capa (Seção 1)
doc.add_picture('logo.jpg', width=Inches(4.5))  # TODO : aqui alterar tambem
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
##################
cpf_formatado = formatar_cpf(cpf)

cabecalhos = ['ITEM', 'DESCRIÇÃO', 'UNID.', 'QUANT.', 'CUSTO UNITÁRIO', 'CUSTO TOTAL']
lista = final_services

##################
add_blank_line(doc, size=24)

title = "PROPOSTA TÉCNICA E ORÇAMENTÁRIA PARA\nPRESTAÇÃO DE SERVIÇOS"
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run(title)
run.font.size = Pt(18)
run.bold = True
run.font.name = 'Times New Roman'

add_blank_line(doc, size=18)

subtitle = f"{objeto_serviço}"
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run(subtitle)
run.font.size = Pt(19)
run.bold = True
run.font.name = 'Times New Roman'

add_blank_line(doc, size=24)

info = f"""
CLIENTE(S): {nome_proponente}
IMÓVEL: {denominacao_imovel}
CIDADE: {municipio}
"""
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run(info)
run.font.size = Pt(14)
run.bold = True
run.font.name = 'Times New Roman'

add_blank_line(doc, size=24)

footer_info = f"""
PALMAS - TO
{mes_atual}, {ano_atual}
"""
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run(footer_info)
run.font.size = Pt(14)
run.bold = True
run.font.name = 'Times New Roman'

# Criar nova seção (Seção 2) para o sumário e conteúdo, onde o rodapé será adicionado
doc.add_section()

# Adicionando o sumário
summary_title = "SUMÁRIO"
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run(summary_title)
run.font.size = Pt(16)
run.bold = True

run.font.name = 'Times New Roman'

add_blank_line(doc, size=18)

# Inicializa a lista com as partes do conteúdo
summary_content = [
    "1. APRESENTAÇÃO\t\n",
    "2. OBJETO DO SERVIÇO SOLICITADO\t\n",
    "3. TRABALHOS A SEREM EXECUTADOS PELA CONTRATADA\t\n"
]

# Adiciona as partes dinâmicas a partir do loop
for i, service in enumerate(final_services):
    summary_content.append(f"   3.{i+1} {service}\t\n")

# Adiciona o restante das partes fixas
summary_content.extend([
    "4. DOCUMENTOS E INFORMAÇÕES DE RESPONSABILIDADE DA CONTRATANTE\t\n",
    "   4.1. Documentos Específicos\t\n",
    "   4.2. Requerente Pessoa Jurídica (P.J)\t\n",
    "5. CUSTO DO SERVIÇO PRESTADO PELA CONTRATADA\t\n",
    "6. FORMA DE PAGAMENTO\t\n",
    "7. PRAZO DE EXECUÇÃO DOS SERVIÇOS\t\n",
    "8. VALIDADE DA PROPOSTA\t\n"
])

# Junta todas as partes em uma única string
summary_content = "".join(summary_content)

# Adicionando o conteúdo do sumário ao parágrafo do documento
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
run = p.add_run(summary_content)
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
p.paragraph_format.line_spacing = Pt(25)

############################################
# Adicionar o rodapé apenas na segunda seção
for section in doc.sections[1:]:
    add_footer(section)

# Adicionar quebra de página antes da apresentação
doc.add_page_break()

# Adicionando a seção de apresentação
salutation = f"A sua senhoria, {ele_ou_ela}. {nome_proponente}"
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
run = p.add_run(salutation)
run.font.size = Pt(12)
run.bold = True
run.font.name = 'Times New Roman'

add_blank_line(doc, size=18)

presentation_title = "1. APRESENTAÇÃO"
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
run = p.add_run(presentation_title)
run.font.size = Pt(14)
run.bold = True
run.font.name = 'Times New Roman'

add_blank_line(doc, size=4)

# Adicionando o conteúdo de apresentação em um único bloco com parágrafos formatados
apresentacao = (
    "\tFundada no ano de 2018, a AGROPASSOS ENGENHARIA iniciou suas atividades voltadas para a área de "
    "treinamentos profissionalizantes para comunidades do setor agropecuário em parceria com o SENAR. "
    "Buscando suprir as necessidades dos clientes de maneira cada vez mais completa, passou a atuar com "
    "consultoria profissional e assistência técnica agropecuária, desenvolvendo metodologias e ferramentas "
    "tecnológicas visando garantir a viabilidade dos empreendimentos.\t\n\n" 
    
    "\tEm 2019, começou atuar também em perícias técnicas judiciais, projetos de licenciamento ambiental, "
    "constituição de servidão administrativa, regularização documental de imóveis urbanos e rurais e "
    "avaliações de imóveis rurais e urbanos.\t\n\n"   
    "\tDessa forma, a AGROPASSOS ENGENHARIA desenvolve projetos e serviços buscando inovações, soluções "
    "tecnológicas e ambientais, que sejam aplicáveis e atendam aos preceitos legais, com foco na excelência "
    "em qualidade, a fim de garantir a satisfação dos nossos clientes, colaboradores e demais envolvidos."
)
# Ajuste o espaçamento entre linhas para evitar espaçamento excessivo

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = p.add_run(apresentacao)
p.paragraph_format.line_spacing = Pt(25)
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

add_blank_line(doc, size=1)
doc.add_page_break()
presentation_title = "2. OBJETO DO SERVIÇO SOLICITADO"
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
run = p.add_run(presentation_title)
run.font.size = Pt(14)
run.bold = True
run.font.name = 'Times New Roman'

add_blank_line(doc, size=4)

corpo_topico_2 = (
    "\tA presente proposta técnica e comercial tem como objetivo a prestação de serviços para a " 
    f"{objeto_serviço} {texto_escopo_opcional} da atividade {atividade} desenvolvida no imóvel {imovel} denominado " 
    f"{denominacao_imovel}, no municipio de {municipio}, de propriedade d{ultima_letra_inscricao} {ele_ou_ela}. {nome_proponente}, inscrit{ultima_letra_inscricao} no CPF nº "
    f"{cpf_formatado}."
)


p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = p.add_run(corpo_topico_2)
p.paragraph_format.line_spacing = Pt(25)
run.font.size = Pt(12)
run.font.name = 'Times New Roman'



add_blank_line(doc, size=15)
doc.add_page_break()
presentation_title = "3. TRABALHOS A SEREM EXECUTADOS PELA CONTRATADA"
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
run = p.add_run(presentation_title)
run.font.size = Pt(14)
run.bold = True
run.font.name = 'Times New Roman'
add_blank_line(doc, size=4)

# Adicionar os serviços e descrições
for i, (service, description) in enumerate(zip(final_services, final_descriptions)):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Justificado para alinhar as bordas esquerda e direita
    
    # Configurar o tab stop
    tab_stops = p.paragraph_format.tab_stops
    tab_stop_position = tab_stops.add_tab_stop(Pt(20))  # Ajuste esse valor conforme necessário
    
    # Adicionar o serviço com negrito
    service_run = p.add_run(f'\t3.{i+1} {service}\t\n')
    service_run.bold = True
    service_run.font.size = Pt(12)
    service_run.font.name = 'Times New Roman'
    # Configurar o tab stop
    
    # Adicionar a descrição
    description_run = p.add_run(f'\t{description}')
    description_run.font.size = Pt(12)
    description_run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = Pt(25)
    

# Adicionar quebra de página 
doc.add_page_break()  
presentation_title = "4.    DOCUMENTOS E INFORMAÇÕES DE RESPONSABILIDADE DA CONTRATANTE"
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = p.add_run(presentation_title)
run.font.size = Pt(14)
run.bold = True
run.font.name = 'Times New Roman'
add_blank_line(doc, size=4)
documentos_texto = (
    
"\t"
)
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
# Configurar o tab stop
tab_stops = p.paragraph_format.tab_stops
tab_stop_position = tab_stops.add_tab_stop(Pt(20))  # Ajuste esse valor conforme necessário

# Formatando "4.1", "4.2", e "4.3" em negrito
run = p.add_run("\t4.1. Documentos Específicos\t")
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'


run = p.add_run("\n\t\t • Certidão de inteiro teor do cartório de registro de imóveis atualizada referente  \t")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.line_spacing = Pt(25)
run = p.add_run("\n\t\t    ao imóvel, ou documento de justa posse;\t")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.line_spacing = Pt(25)
run = p.add_run("\n\t\t • Procuração para representação junto ao órgão ambiental;\t")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.line_spacing = Pt(25)
run = p.add_run("\n\t\t • Comprovante de recolhimento de taxas; e\t")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.line_spacing = Pt(25)
run = p.add_run("\n\t\t • Documentos específicos a serem solicitados pela equipe técnica caso\t")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.line_spacing = Pt(25)
run = p.add_run("\n\t\t    necessário.\t")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.line_spacing = Pt(25)
run = p.add_run("\n\t4.2. Requerente Pessoa Jurídica (P.J)\t")
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.line_spacing = Pt(25)

run = p.add_run("\n\t\t • CNPJ; \t")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.line_spacing = Pt(25)
run = p.add_run("\n\t\t • Contrato Social; \t")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.line_spacing = Pt(25)
run = p.add_run("\n\t\t • Documentos pessoais dos sócios (RG e CPF); e\t")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.line_spacing = Pt(25)
run = p.add_run("\n\t\t • Comprovante de endereço.\t")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.line_spacing = Pt(25)
run = p.add_run("\n\t4.3. Requerente Pessoa Física\t")
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.line_spacing = Pt(25)

run = p.add_run("\n\t\t • CPF; e\t")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.line_spacing = Pt(25)
run = p.add_run("\n\t\t • Comprovante de endereço.\t")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.line_spacing = Pt(25)


add_blank_line(doc, size=20)

presentation_title = (
    "OBS: Demais documentos não relacionados poderão ser solicitados pela equipe técnica da "
    "AGROPASSOS ENGENHARIA para o andamento dos serviços."
)
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = p.add_run(presentation_title)
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.font.bold = True
run.font.italic = True



# Adicionar quebra de página 
doc.add_page_break()  

presentation_title = "5.	CUSTO DO SERVIÇO PRESTADO PELA CONTRATADA"
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = p.add_run(presentation_title)
run.font.size = Pt(14)
run.bold = True
run.font.name = 'Times New Roman'
add_blank_line(doc, size=4)
valor_total_ajustado = criar_tabela_word(cabecalhos,lista,doc)
descricao_topico_5 = (
    f"Conforme apresentado na Tabela, o custo do serviço foi orçado no valor de R$ {valor_total_ajustado:.2f}"
)
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = p.add_run(descricao_topico_5)
p.paragraph_format.line_spacing = Pt(25)
run.font.size = Pt(12)
run.font.name = 'Times New Roman'


obs_opcional_tabela =  pyautogui.prompt("Adicione uma observação caso deseje ")



if obs_opcional_tabela is not None:
    observacao_Tabela = obs_opcional_tabela

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = p.add_run(f'OBS: {observacao_Tabela}')
    p.paragraph_format.line_spacing = Pt(25)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.italic = True
    run.font.bold = True


# Adicionar quebra de página 
doc.add_page_break()  

presentation_title = "6. FORMA DE PAGAMENTO"
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = p.add_run(presentation_title)
run.font.size = Pt(14)
run.bold = True
run.font.name = 'Times New Roman'
add_blank_line(doc, size=4)
print(valor_total_ajustado)
#########################################################
def coletar_dados_e_gerar_paragrafos(valor_total_ajustado, doc):
    root = tk.Tk()
    root.withdraw()  # Esconde a janela principal do Tkinter

    # Pergunta a quantidade de parcelas
    quantidade_parcelas = simpledialog.askinteger(
        "Quantidade de Parcelas", 
        "Quantas parcelas serão?", 
        parent=root,
        minvalue=1,
        initialvalue=1
    )
    
    if quantidade_parcelas is None or quantidade_parcelas <= 0:
        messagebox.showerror("Erro", "Quantidade de parcelas inválida.")
        return
    
    # Calcula o valor padrão de cada parcela
    valor_parcela_padrao = valor_total_ajustado / quantidade_parcelas
    
    # Cria uma nova janela para inserir todos os valores das parcelas
    valores_parcelas = []
    
    def obter_valores():
        try:
            for i in range(quantidade_parcelas):
                valor = float(entry_vars[i].get())
                if valor <= 0:
                    raise ValueError("O valor deve ser positivo.")
                valores_parcelas.append(valor)
            janela_valores.destroy()
        except ValueError:
            messagebox.showerror("Erro", "Por favor, insira valores numéricos válidos e positivos.")
    
    janela_valores = tk.Toplevel(root)
    janela_valores.title("Valores das Parcelas")

    # Definir tamanho mínimo da janela e centralizá-la na tela
    janela_valores.geometry('400x600')  # Ajuste do tamanho da janela para mais espaço
    janela_valores.update_idletasks()
    x = (janela_valores.winfo_screenwidth() // 2) - (janela_valores.winfo_width() // 2)
    y = (janela_valores.winfo_screenheight() // 2) - (janela_valores.winfo_height() // 2)
    janela_valores.geometry(f'+{x}+{y}')

    # Estilo moderno com ttk
    estilo = ttk.Style()
    estilo.configure('TFrame', background='#f0f0f0')
    estilo.configure('TLabel', background='#f0f0f0', font=('Arial', 12))
    estilo.configure('TButton', font=('Arial', 12), padding=10)
    estilo.configure('TEntry', font=('Arial', 12))

    frame = ttk.Frame(janela_valores)
    frame.pack(fill="both", expand=True, padx=20, pady=20)

    instruction_label = ttk.Label(frame, text="Por favor, revise os valores das parcelas:")
    instruction_label.grid(row=0, column=0, columnspan=2, pady=(0, 10))

    entry_vars = []
    for i in range(quantidade_parcelas):
        label = ttk.Label(frame, text=f"Valor da parcela {i+1}:")
        label.grid(row=i+1, column=0, pady=5, sticky="e")
        entry_var = tk.StringVar(value=f"{valor_parcela_padrao:.2f}")  # Preenche com o valor padrão
        entry_vars.append(entry_var)
        entry = ttk.Entry(frame, textvariable=entry_var)
        entry.grid(row=i+1, column=1, pady=5, sticky="w")

    submit_button = ttk.Button(frame, text="Confirmar", command=obter_valores)
    submit_button.grid(row=quantidade_parcelas+1, column=0, columnspan=2, pady=20)

    root.wait_window(janela_valores)

    if not valores_parcelas:
        return

    # Processa cada parcela inserida
    for i, parcelas in enumerate(valores_parcelas):
        percentual = calcular_percentual(parcelas, valor_total_ajustado)

        descricao_parcela = (
            f'   • {percentual:.1f}% ({i+1}ª parcela) a ser pago na assinatura '
            f'(aqui vai o conteúdo, vai depender das opções) (R$ {parcelas:.2f})'
        )

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        run = p.add_run(descricao_parcela)
        p.paragraph_format.line_spacing = Pt(25)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

coletar_dados_e_gerar_paragrafos(valor_total_ajustado, doc)
#########################################################
add_blank_line(doc, size=6)
presentation_title = "7.	PRAZO DE EXECUÇÃO DOS SERVIÇOS"
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = p.add_run(presentation_title)
run.font.size = Pt(14)
run.bold = True
run.font.name = 'Times New Roman'
add_blank_line(doc, size=4)
escopo_do_7 = (
    f"\tO serviço será entregue e protocolizado junto ao órgão ambiental, em até {prazo} dias, desde que cumpridas todas as exigências a cargo da contratante."
)
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = p.add_run(escopo_do_7)
p.paragraph_format.line_spacing = Pt(25)
run.font.size = Pt(12)
run.font.name = 'Times New Roman'  

# Adicionar quebra de página 
doc.add_page_break()  

add_blank_line(doc, size=10)
presentation_title = "8.	VALIDADE DA PROPOSTA"
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = p.add_run(presentation_title)
run.font.size = Pt(14)
run.bold = True
run.font.name = 'Times New Roman'
add_blank_line(doc, size=4)
escopo_do_8 = (
"\tEsta proposta tem validade de 15 (quinze) dias a contar de sua apresentação podendo este prazo ser prorrogado conforme anuência das partes.\t\n"
"\tPor fim, nos colocamos a disposição para esclarecimento de qualquer dúvida referente à proposta apresentada."
)
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = p.add_run(escopo_do_8)
p.paragraph_format.line_spacing = Pt(25)
run.font.size = Pt(12)
run.font.name = 'Times New Roman' 
add_blank_line(doc, size=6)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
run = p.add_run(f'Palmas {dia_atual} de {mes_atual} de {ano_atual}')
p.paragraph_format.line_spacing = Pt(25)
run.font.size = Pt(12)
run.font.name = 'Times New Roman' 
#add um espaço minimo entre palmas e o resto do texto
add_blank_line(doc, size=4)

doc.add_picture('agro_passos_assinatura.png', width=Inches(2.5))  # TODO : aqui alterar tambem
assinatura_agropassos = doc.paragraphs[-1]
assinatura_agropassos.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
add_blank_line(doc, size=3)
# Criação da tabela com 1 linha e 2 colunas
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Inserção da primeira imagem na primeira célula
cell_1 = table.cell(0, 0)
cell_1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
paragraph_1 = cell_1.paragraphs[0]
run_1 = paragraph_1.add_run()
run_1.add_picture('assinatura_arthur.png', width=Inches(2.5))#TODO : aqui alterar pra planilha do user

# Inserção da segunda imagem na segunda célula
cell_2 = table.cell(0, 1)
cell_2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
paragraph_2 = cell_2.paragraphs[0]
run_2 = paragraph_2.add_run()
run_2.add_picture('assinatura_luan.png', width=Inches(2.5))#TODO : aqui alterar pra planilha do user
doc.save('proposta_tecnica_orcamentaria.docx')# TODO : aqui vai depender de como voce quer salvar em uma pasta especifica

# Abrindo o arquivo Word após salvá-lo
os.startfile('proposta_tecnica_orcamentaria.docx') # TODO : aqui vai depender de como voce quer salvar em uma pasta especifica